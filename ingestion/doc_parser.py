import os
from pathlib import Path
from typing import Optional


try:
    import pypdf
except ImportError:
    pypdf = None


try:
    import docx
except ImportError:
    docx = None

def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from a PDF document page by page."""
    if not pypdf:
        raise RuntimeError("pypdf is not installed.")


    reader = pypdf.PdfReader(str(file_path))
    pages_text = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text and text.strip():
            pages_text.append(f"--- [Page {page_num}] ---\n{text.strip()}")
    return "\n\n".join(pages_text)




def extract_text_from_docx(file_path: Path) -> str:
    """Extract paragraphs and table contents from a Microsoft Word (.docx) document."""
    if not docx:
        raise RuntimeError("python-docx is not installed.")


    doc = docx.Document(str(file_path))
    sections = []
    for para in doc.paragraphs:
        if para.text.strip():
            sections.append(para.text.strip())


    # Also extract text from tables if present
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(row_cells))
        if table_rows:
            sections.append("\n".join(table_rows))


    return "\n\n".join(sections)




def extract_text(file_path: Path) -> str:
    """Unified text extractor for PDF, DOCX, MD, TXT files."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in (".md", ".txt", ".json", ".csv"):
        return file_path.read_text(encoding="utf-8", errors="ignore")
    else:
        # Fallback reading as plain text
        try:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            return f"Could not parse file {file_path.name}: {str(e)}"
